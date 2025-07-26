import os
import re
import cv2
import json
import random
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches
from typing import Dict, Tuple, List, Optional
from config import Config
import logging

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Enhanced document generation with simplified duplicate handling"""
    
    @staticmethod
    def create_complete_document(id_info: Dict[str, str], 
                               front_image: np.ndarray, 
                               back_image: np.ndarray,
                               qr_content: str,
                               overwrite: bool = False,
                               existing_file_info: Optional[Dict[str, str]] = None) -> Tuple[str, List[str]]:
        """Create comprehensive ID document with simplified overwrite support"""
        old_files_to_delete = []
        
        try:
            doc = Document()
            
            # Title and metadata
            DocumentGenerator._add_header(doc, id_info)
            
            # Information table
            DocumentGenerator._add_info_table(doc, id_info)
            
            # ALWAYS create new files with unique names
            front_path, back_path = DocumentGenerator._save_images_safe(
                front_image, back_image, id_info
            )
            
            # Photos section
            DocumentGenerator._add_photos_section(doc, front_path, back_path)
            
            # ALWAYS create new document with unique name
            filename = DocumentGenerator._generate_filename_safe(id_info)
            doc.save(filename)
            
            logger.info(f"Document saved successfully: {filename}")
            
            # If overwrite requested, prepare list of old files to delete
            if overwrite and existing_file_info:
                if 'json_path' in existing_file_info and os.path.exists(existing_file_info['json_path']):
                    old_files_to_delete.append(existing_file_info['json_path'])
                
                if 'front_image' in existing_file_info and os.path.exists(existing_file_info['front_image']):
                    old_files_to_delete.append(existing_file_info['front_image'])
                
                if 'back_image' in existing_file_info and os.path.exists(existing_file_info['back_image']):
                    old_files_to_delete.append(existing_file_info['back_image'])
                
                if 'word_path' in existing_file_info and os.path.exists(existing_file_info['word_path']):
                    old_files_to_delete.append(existing_file_info['word_path'])
            
            return os.path.basename(filename), old_files_to_delete
            
        except Exception as e:
            logger.error(f"Failed to create document: {e}")
            raise
    
    @staticmethod
    def _save_images_safe(front_image: np.ndarray, back_image: np.ndarray, 
                         id_info: Dict[str, str]) -> tuple:
        """Save images with robust error handling"""
        front_path = None
        back_path = None
        json_path = None
        
        try:
            # Ensure directories exist
            os.makedirs(Config.SCAN_DIR, exist_ok=True)
            logger.info(f"Scan directory ensured: {Config.SCAN_DIR}")
            
            # Generate safe base name
            safe_name = re.sub(r'[^\w\s-]', '', id_info.get("Họ và tên", "Unknown"))
            safe_name = re.sub(r'[-\s]+', '_', safe_name)
            if not safe_name:
                safe_name = "Unknown"
            
            # Simple timestamp to avoid issues
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            random_suffix = f"{random.randint(1000, 9999)}"
            
            # Shorter filename to avoid path length issues
            base_filename = f"ID_{timestamp}_{random_suffix}"
            
            logger.info(f"Base filename: {base_filename}")
            
            # Generate file paths
            front_filename = f"{base_filename}_F.jpg"
            back_filename = f"{base_filename}_B.jpg"
            json_filename = f"{base_filename}.json"
            
            front_path = os.path.join(Config.SCAN_DIR, front_filename)
            back_path = os.path.join(Config.SCAN_DIR, back_filename)
            json_path = os.path.join(Config.SCAN_DIR, json_filename)
            
            # Save front image
            success = cv2.imwrite(front_path, front_image, [cv2.IMWRITE_JPEG_QUALITY, 95])
            if not success or not os.path.exists(front_path):
                raise Exception("Failed to save front image")
            
            # Save back image
            success = cv2.imwrite(back_path, back_image, [cv2.IMWRITE_JPEG_QUALITY, 95])
            if not success or not os.path.exists(back_path):
                raise Exception("Failed to save back image")
            
            # Save JSON data
            json_data = id_info.copy()
            json_data['_filename_base'] = base_filename
            json_data['_saved_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"✅ All files saved successfully!")
            return front_path, back_path
            
        except Exception as e:
            logger.error(f"Failed to save images: {str(e)}")
            raise Exception(f"Image save failed: {str(e)}")
    
    @staticmethod
    def _generate_filename_safe(id_info: Dict[str, str]) -> str:
        """Generate safe filename - ALWAYS creates unique filename"""
        safe_name = re.sub(r'[^\w\s-]', '', id_info.get("Họ và tên", "Unknown"))
        safe_name = re.sub(r'[-\s]+', '_', safe_name)[:20]  # Limit length
        
        if not safe_name:
            safe_name = "Unknown"
        
        # Simpler timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        random_suffix = f"{random.randint(1000, 9999)}"
        
        # Ensure output directory exists
        os.makedirs(Config.OUTPUT_DIR, exist_ok=True)
        
        # Shorter filename
        filename = f"CCCD_{safe_name}_{timestamp}_{random_suffix}.docx"
        full_path = os.path.join(Config.OUTPUT_DIR, filename)
        
        # Ensure unique
        counter = 1
        while os.path.exists(full_path) and counter < 100:
            filename = f"CCCD_{safe_name}_{timestamp}_{random_suffix}_{counter}.docx"
            full_path = os.path.join(Config.OUTPUT_DIR, filename)
            counter += 1
        
        return full_path
    
    @staticmethod
    def _add_header(doc: Document, id_info: Dict[str, str]):
        """Add document header"""
        title = doc.add_heading('HỒ SƠ CCCD/CMND HOÀN CHỈNH', level=0)
        title.alignment = 1
        
        current_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        doc.add_paragraph(f"Thời gian tạo: {current_time}")
        separator_para = doc.add_paragraph("═" * 50)
        separator_para.alignment = 1
    
    @staticmethod
    def _add_info_table(doc: Document, id_info: Dict[str, str]):
        """Add information table"""
        doc.add_heading('THÔNG TIN CÁ NHÂN (QR CODE)', level=2)
        
        table = doc.add_table(rows=len(id_info), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(id_info.items()):
            table.cell(i, 0).text = f"{label}:"
            table.cell(i, 1).text = value
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    @staticmethod
    def _add_photos_section(doc: Document, front_path: str, back_path: str):
        """Add photos section"""
        doc.add_heading('ẢNH CCCD', level=2)
        
        # Create table for side-by-side images
        photo_table = doc.add_table(rows=2, cols=2)
        photo_table.autofit = False
        
        # Set column widths
        for cell in photo_table.columns[0].cells:
            cell.width = Inches(3.2)
        for cell in photo_table.columns[1].cells:
            cell.width = Inches(3.2)
        
        # Add headers
        photo_table.cell(0, 0).text = 'MẶT TRƯỚC'
        photo_table.cell(0, 1).text = 'MẶT SAU'
        
        # Center align headers
        for i in range(2):
            photo_table.cell(0, i).paragraphs[0].alignment = 1
            photo_table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Add images to table
        for i, path in enumerate([front_path, back_path]):
            if os.path.exists(path):
                paragraph = photo_table.cell(1, i).paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(path, width=Inches(3.0))
                paragraph.alignment = 1
        
        # Add verification note
        separator_para = doc.add_paragraph("═" * 50)
        separator_para.alignment = 1
        
        current_date = datetime.now()
        day = current_date.day
        month = current_date.month
        year = current_date.year
        
        verification_para1 = doc.add_paragraph("Đã đối chiếu bản chính.")
        verification_para1.alignment = 1
        verification_para1.runs[0].bold = True
        
        date_para = doc.add_paragraph(f"Bạc Liêu, ngày {day}, tháng {month}, năm {year}.")
        date_para.alignment = 1
        
        staff_para = doc.add_paragraph("Giao dịch viên.")
        staff_para.alignment = 1
        staff_para.runs[0].italic = True
